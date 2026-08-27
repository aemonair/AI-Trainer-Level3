#!/usr/bin/env python3
"""
提取第3部分-理论知识复习题中的所有题目并制作Anki卡片
支持：判断题、单选题、多选题
"""
import docx
import re
import os
import csv

os.chdir('/Users/air/Downloads/GUIDE_AI_3/人工智能训练师_3级_sucai')

print("="*60)
print("提取理论知识复习题 制作Anki卡片")
print("="*60)

# 读取docx文件
doc = docx.Document('4-04-05-05_3_20250701/第3部分-人工智能训练师_3级_理论知识复习题.docx')

# 1. 查找所有题型标题
section_headers = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if re.match(r'^[一二三四五六七八九十]+、', text):
        section_headers.append((i, text))

print(f"\n找到 {len(section_headers)} 个题型部分:")
for idx, (line_num, header) in enumerate(section_headers):
    print(f"  {idx+1}. 段落{line_num}: {header}")

# 2. 按题型分类提取题目
questions = {
    '判断题': [],
    '单选题': [],
    '多选题': []
}

current_section = None
current_question = None
question_buffer = []

for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    
    # 检测题型标题
    for idx, (line_num, header) in enumerate(section_headers):
        if header in text:
            current_section = header
            if '判断' in header:
                current_section = '判断题'
            elif '单选' in header:
                current_section = '单选题'
            elif '多选' in header:
                current_section = '多选题'
            break
    
    # 根据题型特征提取题目
    if current_section == '判断题':
        if text.startswith('（    ）') or text.startswith('(    )'):
            questions['判断题'].append(text)
    
    elif current_section == '单选题':
        # 单选题通常有题号+选项
        if re.match(r'^\d+[\.\．]', text) or re.match(r'^[（(]\s*[）)]', text):
            if current_question:
                questions['单选题'].append('\n'.join(question_buffer))
                question_buffer = []
            current_question = text
            question_buffer.append(text)
        elif text and ('A．' in text or 'A.' in text or 'B．' in text or 'B.' in text):
            question_buffer.append(text)
    
    elif current_section == '多选题':
        if re.match(r'^\d+[\.\．]', text) or re.match(r'^[（(]\s*[）)]', text):
            if current_question:
                questions['多选题'].append('\n'.join(question_buffer))
                question_buffer = []
            current_question = text
            question_buffer.append(text)
        elif text and ('A．' in text or 'A.' in text):
            question_buffer.append(text)

# 保存最后一个题目
if question_buffer:
    if current_section == '单选题':
        questions['单选题'].append('\n'.join(question_buffer))
    elif current_section == '多选题':
        questions['多选题'].append('\n'.join(question_buffer))

# 3. 统计结果
print("\n" + "="*60)
print("题目数量统计")
print("="*60)
total = 0
for qtype, qlist in questions.items():
    count = len(qlist)
    total += count
    print(f"{qtype}: {count} 题")
print(f"总计: {total} 题")

# 4. 生成Anki卡片
print("\n" + "="*60)
print("生成Anki卡片...")
print("="*60)

os.makedirs('anki_cards', exist_ok=True)

# 为每种题型生成单独的CSV
for qtype, qlist in questions.items():
    if not qlist:
        continue
    
    filename = f'anki_cards/理论知识_{qtype}.csv'
    with open(filename, 'w', encoding='utf-8', newline='') as f:
        writer = csv.writer(f, delimiter=';')
        writer.writerow(['正面', '背面'])
        
        for i, q in enumerate(qlist):
            # 正面：题目
            front = f"【{qtype}】{q}"
            # 背面：答案（暂时为空，需要手动填写）
            back = f"【答案】\n\n（请手动填写答案）\n\n【题目原文】\n{q}"
            writer.writerow([front, back])
    
    print(f"✅ {qtype}: {len(qlist)}题 → {filename}")

# 5. 生成合并版
print("\n生成合并版Anki卡片...")
all_filename = 'anki_cards/理论知识_全部题目.csv'
with open(all_filename, 'w', encoding='utf-8', newline='') as f:
    writer = csv.writer(f, delimiter=';')
    writer.writerow(['正面', '背面'])
    
    for qtype, qlist in questions.items():
        for i, q in enumerate(qlist):
            front = f"【{qtype}】{q}"
            back = f"【答案】\n\n（请手动填写答案）\n\n【题目原文】\n{q}"
            writer.writerow([front, back])

print(f"✅ 合并版: {total}题 → {all_filename}")

# 6. 显示示例
print("\n" + "="*60)
print("判断题示例（前3题）:")
print("="*60)
for i, q in enumerate(questions['判断题'][:3]):
    print(f"\n{i+1}. {q[:100]}...")

print("\n" + "="*60)
print("单选题示例（前2题）:")
print("="*60)
for i, q in enumerate(questions['单选题'][:2]):
    print(f"\n{i+1}. {q[:150]}...")

print("\n" + "="*60)
print("多选题示例（前2题）:")
print("="*60)
for i, q in enumerate(questions['多选题'][:2]):
    print(f"\n{i+1}. {q[:150]}...")

print("\n" + "="*60)
print("✅ 完成！")
print("="*60)
print("\n📁 生成的文件:")
print("  - anki_cards/理论知识_判断题.csv")
print("  - anki_cards/理论知识_单选题.csv")
print("  - anki_cards/理论知识_多选题.csv")
print("  - anki_cards/理论知识_全部题目.csv")
print("\n💡 提示：导入Anki后需要手动填写答案")