# ============================================================
# 3.1.X 智能产品数据分析通用模板
# 适用于：3.1.1 ~ 3.1.5 所有数据分析题目
# 使用方法：
#   1. 修改 FILE_PATH 为实际Excel文件路径
#   2. 根据具体题目修改列名
#   3. 运行脚本生成分析报告
# ============================================================

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Inches
import os

# ============================================================
# 配置区域 - 根据具体题目修改以下参数
# ============================================================

# 文件路径配置
FILE_PATH = 'template/3.1.5/智能家居环境控制系统数据集.xlsx'  # 修改为实际路径
OUTPUT_DIR = 'output'  # 输出目录

# 列名配置 - 根据实际Excel列名修改
CONFIG = {
    '3.1.1': {  # 智能音箱
        'time_col': '时间',
        'function_col': '功能调用类型',
        'response_time_col': '响应时间',
        'time_periods': ['06:00-12:00', '12:00-18:00', '18:00-24:00'],
    },
    '3.1.2': {  # 智能照明系统
        'time_col': '时间',
        'brightness_col': '光线亮度',
        'color_temp_col': '色温',
        'scene_col': '智能场景',
        'response_time_col': '响应时间',
        'time_periods': ['06:00-12:00', '12:00-18:00', '18:00-24:00'],
    },
    '3.1.3': {  # 智能健康手环
        'time_col': '时间',
        'steps_col': '步数',
        'heart_rate_col': '心率',
        'sleep_col': '睡眠时长',
        'sync_delay_col': '同步延迟',
        'time_periods': ['06:00-08:00', '17:00-20:00', '其余时段'],
    },
    '3.1.4': {  # 智能健康监测系统
        'time_col': '时间',
        'systolic_col': '收缩压',
        'diastolic_col': '舒张压',
        'blood_sugar_col': '血糖值',
        'body_fat_col': '体脂值',
        'response_time_col': '响应时间',
    },
    '3.1.5': {  # 智能家居环境控制系统
        'time_col': '时间',
        'temp_col': '温度',
        'humidity_col': '湿度',
        'light_col': '光照',
        'response_time_col': '响应时间',
        'energy_col': '能源消耗',
        'time_periods': ['06:00-12:00', '13:00-18:00', '19:00-05:00'],
    }
}

# 当前分析的题目编号
CHAPTER = '3.1.5'  # 修改为当前题目：3.1.1 ~ 3.1.5


# ============================================================
# 工具函数
# ============================================================

def read_excel(file_path):
    """读取Excel文件"""
    print(f"正在读取文件: {file_path}")
    df = pd.read_excel(file_path)
    print(f"数据加载完成，共 {len(df)} 行，{len(df.columns)} 列")
    print(f"列名: {list(df.columns)}")
    return df


def parse_time_column(df, time_col):
    """解析时间列，提取小时"""
    df['hour'] = pd.to_datetime(df[time_col]).dt.hour
    return df


def get_time_period(hour, periods_type='default'):
    """根据小时返回时间段"""
    if periods_type == '3.1.5':
        # 3.1.5 特殊时间段划分
        if 6 <= hour < 12:
            return '06:00 - 12:00'
        elif 12 <= hour < 18:
            return '13:00 - 18:00'
        else:
            return '19:00 - 05:00'
    elif periods_type == '3.1.3':
        # 3.1.3 特殊时间段划分
        if 6 <= hour < 8:
            return '06:00 - 08:00'
        elif 17 <= hour < 20:
            return '17:00 - 20:00'
        else:
            return '其余时段'
    else:
        # 默认时间段划分
        if 6 <= hour < 12:
            return '06:00 - 12:00'
        elif 12 <= hour < 18:
            return '12:00 - 18:00'
        else:
            return '18:00 - 24:00'


def calculate_stats_by_period(df, value_cols, time_col, periods_type='default'):
    """按时间段计算统计值"""
    df = parse_time_column(df, time_col)
    df['period'] = df['hour'].apply(lambda h: get_time_period(h, periods_type))
    
    stats = df.groupby('period')[value_cols].mean()
    print("\n=== 按时间段统计 ===")
    print(stats)
    return stats


def calculate_response_time_analysis(df, response_col, function_col=None):
    """响应时间分析"""
    avg_response = df[response_col].mean()
    print(f"\n平均响应时间: {avg_response:.2f}")
    
    if function_col:
        response_by_func = df.groupby(function_col)[response_col].mean()
        print("\n各功能平均响应时间:")
        print(response_by_func)
        
        # 找出响应时间较长/适中/较短的功能
        sorted_funcs = response_by_func.sort_values(ascending=False)
        print(f"\n响应时间较长的功能: {sorted_funcs.index[0]}")
        print(f"响应时间较短的功能: {sorted_funcs.index[-1]}")
        
        return avg_response, response_by_func
    
    return avg_response, None


def calculate_frequency_analysis(df, category_col):
    """频率分析"""
    freq = df[category_col].value_counts()
    print(f"\n=== {category_col} 使用频率 ===")
    print(freq)
    return freq


def create_word_report(chapter, analysis_data, output_dir='output'):
    """生成Word报告"""
    os.makedirs(output_dir, exist_ok=True)
    
    doc = Document()
    
    # 添加标题
    doc.add_heading(f'{chapter} 分析报告', level=1)
    
    # 添加分析内容
    for section_title, content in analysis_data.items():
        doc.add_heading(section_title, level=2)
        if isinstance(content, pd.DataFrame):
            # 添加表格
            table = doc.add_table(rows=len(content)+1, cols=len(content.columns)+1)
            # 表头
            for j, col in enumerate(content.columns):
                table.rows[0].cells[j].text = col
            # 数据行
            for i, (idx, row) in enumerate(content.iterrows()):
                table.rows[i+1].cells[0].text = str(idx)
                for j, val in enumerate(row):
                    table.rows[i+1].cells[j+1].text = str(round(val, 2))
        else:
            doc.add_paragraph(str(content))
    
    # 保存文件
    output_path = os.path.join(output_dir, f'{chapter}-1.docx')
    doc.save(output_path)
    print(f"\n报告已保存至: {output_path}")
    return output_path


# ============================================================
# 各题目专用分析函数
# ============================================================

def analyze_3_1_1(df):
    """3.1.1 智能音箱产品分析"""
    config = CONFIG['3.1.1']
    
    print("\n" + "="*50)
    print("3.1.1 智能音箱产品数据分析")
    print("="*50)
    
    # 1. 用户使用习惯 - 最常被使用的功能
    print("\n【用户使用习惯】")
    func_freq = calculate_frequency_analysis(df, config['function_col'])
    top_funcs = func_freq.head(3)
    print(f"\n最常被使用的功能: {list(top_funcs.index)}")
    
    # 2. 功能使用频率
    print("\n【功能使用频率】")
    print(f"最受欢迎的功能: {func_freq.index[0]}")
    print(f"较少使用的功能: {list(func_freq.index[-2:])}")
    
    # 3. 响应时间分析
    print("\n【响应时间分析】")
    avg_resp, resp_by_func = calculate_response_time_analysis(
        df, config['response_time_col'], config['function_col']
    )
    
    # 生成分析报告
    analysis_data = {
        '用户使用习惯': f'最常被使用的功能: {", ".join(top_funcs.index)}',
        '功能使用频率': func_freq,
        '响应时间分析': resp_by_func if resp_by_func is not None else f'平均响应时间: {avg_resp:.2f}'
    }
    
    create_word_report('3.1.1', analysis_data)


def analyze_3_1_2(df):
    """3.1.2 智能照明系统分析"""
    config = CONFIG['3.1.2']
    
    print("\n" + "="*50)
    print("3.1.2 智能照明系统数据分析")
    print("="*50)
    
    # 1. 用户使用习惯 - 按时间段统计
    print("\n【用户使用习惯】")
    stats = calculate_stats_by_period(
        df, 
        [config['brightness_col'], config['color_temp_col']], 
        config['time_col']
    )
    
    # 2. 智能场景使用频率
    print("\n【智能场景使用频率】")
    scene_freq = calculate_frequency_analysis(df, config['scene_col'])
    
    # 3. 响应时间分析
    print("\n【响应时间分析】")
    avg_resp = calculate_response_time_analysis(df, config['response_time_col'])[0]
    
    # 生成分析报告
    analysis_data = {
        '用户使用习惯': stats,
        '智能场景使用频率': scene_freq,
        '响应时间分析': f'平均响应时间: {avg_resp:.2f}'
    }
    
    create_word_report('3.1.2', analysis_data)


def analyze_3_1_3(df):
    """3.1.3 智能健康手环分析"""
    config = CONFIG['3.1.3']
    
    print("\n" + "="*50)
    print("3.1.3 智能健康手环数据分析")
    print("="*50)
    
    # 1. 用户活动模式
    print("\n【用户活动模式】")
    steps_stats = calculate_stats_by_period(
        df, [config['steps_col']], config['time_col'], '3.1.3'
    )
    
    # 2. 健康指标关注度
    print("\n【健康指标关注度】")
    print(f"平均步数: {df[config['steps_col']].mean():.2f}")
    print(f"平均心率: {df[config['heart_rate_col']].mean():.2f}")
    print(f"平均睡眠时长: {df[config['sleep_col']].mean():.2f}")
    
    # 3. 数据同步性能
    print("\n【数据同步性能】")
    avg_delay = df[config['sync_delay_col']].mean()
    print(f"平均延迟时间: {avg_delay:.2f}秒")
    
    # 生成分析报告
    analysis_data = {
        '用户活动模式': steps_stats,
        '健康指标关注度': f'步数: {df[config["steps_col"]].mean():.2f}\n心率: {df[config["heart_rate_col"]].mean():.2f}\n睡眠时长: {df[config["sleep_col"]].mean():.2f}',
        '数据同步性能': f'平均延迟: {avg_delay:.2f}秒'
    }
    
    create_word_report('3.1.3', analysis_data)


def analyze_3_1_4(df):
    """3.1.4 智能健康监测系统分析"""
    config = CONFIG['3.1.4']
    
    print("\n" + "="*50)
    print("3.1.4 智能健康监测系统数据分析")
    print("="*50)
    
    # 1. 用户活动周期
    print("\n【用户活动周期】")
    df = parse_time_column(df, config['time_col'])
    
    # 血压趋势
    morning_bp = df[df['hour'].between(6, 8)][config['systolic_col']].mean()
    print(f"早晨平均收缩压: {morning_bp:.2f}")
    
    # 血糖趋势
    meal_times = df[df['hour'].between(7, 9) | df['hour'].between(12, 14) | df['hour'].between(18, 20)]
    meal_blood_sugar = meal_times[config['blood_sugar_col']].mean()
    print(f"进餐后平均血糖: {meal_blood_sugar:.2f}")
    
    # 2. 健康指标偏好度
    print("\n【健康指标偏好度】")
    print(f"血压监测数据量: {len(df[df[config['systolic_col']].notna()])}")
    print(f"血糖检测数据量: {len(df[df[config['blood_sugar_col']].notna()])}")
    print(f"体脂分析数据量: {len(df[df[config['body_fat_col']].notna()])}")
    
    # 3. 系统响应与准确性
    print("\n【系统响应与准确性】")
    avg_resp = calculate_response_time_analysis(df, config['response_time_col'])[0]
    
    # 生成分析报告
    analysis_data = {
        '用户活动周期': f'早晨血压趋势、进餐后血糖趋势',
        '健康指标偏好度': f'血压: {len(df[df[config["systolic_col"]].notna()])}条\n血糖: {len(df[df[config["blood_sugar_col"]].notna()])}条\n体脂: {len(df[df[config["body_fat_col"]].notna()])}条',
        '系统响应与准确性': f'平均响应时间: {avg_resp:.2f}'
    }
    
    create_word_report('3.1.4', analysis_data)


def analyze_3_1_5(df):
    """3.1.5 智能家居环境控制系统分析"""
    config = CONFIG['3.1.5']
    
    print("\n" + "="*50)
    print("3.1.5 智能家居环境控制系统数据分析")
    print("="*50)
    
    # 1. 用户环境偏好
    print("\n【用户环境偏好】")
    stats = calculate_stats_by_period(
        df, 
        [config['temp_col'], config['humidity_col'], config['light_col']], 
        config['time_col'],
        '3.1.5'
    )
    
    # 2. 系统响应时间
    print("\n【系统响应时间】")
    avg_resp = calculate_response_time_analysis(df, config['response_time_col'])[0]
    print(f"平均响应时间: {avg_resp:.2f}")
    print("影响因素: 网络延迟、系统处理能力")
    
    # 3. 能源消耗分析
    print("\n【能源消耗分析】")
    avg_energy = df[config['energy_col']].mean()
    print(f"平均能源消耗: {avg_energy:.2f}")
    print("节能潜力: 优化算法减少不必要操作、预测模型避免频繁启停")
    
    # 生成分析报告
    analysis_data = {
        '用户环境偏好': stats,
        '系统响应时间': f'平均响应时间: {avg_resp:.2f}\n影响因素: 网络延迟、系统处理能力',
        '能源消耗分析': f'平均能源消耗: {avg_energy:.2f}\n节能潜力: 优化算法、预测模型'
    }
    
    create_word_report('3.1.5', analysis_data)


# ============================================================
# 主函数
# ============================================================

def main():
    """主函数 - 根据章节号调用相应的分析函数"""
    
    # 读取数据
    df = read_excel(FILE_PATH)
    
    # 根据章节调用相应分析函数
    if CHAPTER == '3.1.1':
        analyze_3_1_1(df)
    elif CHAPTER == '3.1.2':
        analyze_3_1_2(df)
    elif CHAPTER == '3.1.3':
        analyze_3_1_3(df)
    elif CHAPTER == '3.1.4':
        analyze_3_1_4(df)
    elif CHAPTER == '3.1.5':
        analyze_3_1_5(df)
    else:
        print(f"错误: 不支持的章节 {CHAPTER}")
        print("支持的章节: 3.1.1, 3.1.2, 3.1.3, 3.1.4, 3.1.5")


if __name__ == '__main__':
    main()